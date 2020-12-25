from docx import Document
from docx.shared import Inches

document = Document()

# Profile picture
document.add_picture('adel.png', width=Inches(2.0))

# Profile details
name = input('What is your name: ')
age = input('How old are you: ')
phone_number = input('What is your phone number: ')
email = input('What is your email: ')

document.add_paragraph(f"Name: {name}")
document.add_paragraph(f"Age: {age}")
document.add_paragraph(f"Phone Number: {phone_number}")
document.add_paragraph(f"Email: {email}")

# About
document.add_heading('About')
document.add_paragraph(input('Tell about yourself: '))

# Experience
document.add_heading('Work Experience')
p = document.add_paragraph()


def add_experience():
    company = input('Enter company: ')
    from_date = input('From Date: ')
    to_date = input('To Date: ')
    p.add_run(company + ' ').bold = True
    p.add_run(from_date + '-' + to_date + '\n').italic = True
    experience_details = input(f"Describe your experience at '{company}' : ")
    p.add_run("Description: \n" + experience_details)


add_experience()

# Other experiences
while True:
    has_more_experiences = input('Do you have more experiences? (Yes) or (No): ')
    if has_more_experiences == "Yes" or has_more_experiences == "yes":
        p = document.add_paragraph()
        add_experience()
    else:
        break

# Skills
document.add_heading('Skills')


def add_skills():
    skill = input('Enter skill: ')
    p = document.add_paragraph(skill)
    p.style = 'List Bullet'


add_skills()

while True:
    has_more_skills = input('Do you have more skills? (Yes) or (No): ')
    if has_more_skills == 'Yes' or has_more_skills == 'yes':
        add_skills()
    else:
        break

# Footer
section = document.sections[0]
footer = section.footer
paragraph = footer.paragraphs[0]
paragraph.text = f"{name}"
paragraph.alignment = 1

document.save('cv.docx')
